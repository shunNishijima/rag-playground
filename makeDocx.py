import os
import glob
import xml.etree.ElementTree as ET
from docx import Document

input_dir = 'data/xml'     # XMLファイルが置かれているディレクトリ
output_dir = 'data/docx'   # DOCXファイルの出力先
os.makedirs(output_dir, exist_ok=True)

def xml_to_docx(xml_path, docx_path):
    try:
        tree = ET.parse(xml_path)
        root = tree.getroot()
        doc = Document()

        # タイトル
        law_title = root.findtext(".//LawTitle")
        if law_title:
            doc.add_heading(law_title.strip(), level=1)

        # 本文処理
        for article in root.findall(".//Article"):
            title = article.findtext("ArticleTitle")
            caption = article.findtext("ArticleCaption")
            if title:
                doc.add_heading(f"{title} {caption or ''}".strip(), level=2)

            for paragraph in article.findall("Paragraph"):
                # Paragraph > Sentence
                for sentence in paragraph.findall("Sentence"):
                    if sentence.text:
                        doc.add_paragraph(sentence.text.strip())

                # Paragraph > Item > Sentence
                for item in paragraph.findall("Item"):
                    for sentence in item.findall("Sentence"):
                        if sentence.text:
                            doc.add_paragraph(sentence.text.strip())

        doc.save(docx_path)
        print(f"✅ Saved DOCX: {os.path.basename(docx_path)}")

    except ET.ParseError:
        print(f"⚠️ XML parse error: {xml_path}")
    except Exception as e:
        print(f"⚠️ Error converting {xml_path}: {e}")

# 実行
for filepath in glob.glob(os.path.join(input_dir, '*.xml')):
    filename = os.path.basename(filepath).replace('.xml', '.docx')
    docx_path = os.path.join(output_dir, filename)
    xml_to_docx(filepath, docx_path)
